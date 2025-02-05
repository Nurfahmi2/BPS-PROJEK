from flask import Flask, render_template, request, send_file, session
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import io
import os
import subprocess
from datetime import datetime

app = Flask(__name__)
app.secret_key = "secret_key"  
TEMPLATE_FOLDER = 'templates'

# Fungsi untuk mengubah font teks di dokumen
def change_font(doc):
    for para in doc.paragraphs:
        for run in para.runs:
            run.font.name = "Arial"
            run.font.size = Pt(11)
            r = run._element
            r.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')

# Fungsi untuk mengubah tabel ke font Arial 11
def change_table_fonts(table):
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.name = "Arial"
                    run.font.size = Pt(11)
                    r = run._element
                    r.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')

# Fungsi untuk memformat tanggal ke bahasa Indonesia
def format_tanggal(tanggal_str):
    bulan_inggris_ke_indonesia = {
        "January": "Januari", "February": "Februari", "March": "Maret",
        "April": "April", "May": "Mei", "June": "Juni",
        "July": "Juli", "August": "Agustus", "September": "September",
        "October": "Oktober", "November": "November", "December": "Desember"
    }
    try:
        dt = datetime.strptime(tanggal_str, '%Y-%m-%d')
        tanggal_format = dt.strftime('%d %B %Y')
        for eng, ind in bulan_inggris_ke_indonesia.items():
            tanggal_format = tanggal_format.replace(eng, ind)
        return tanggal_format
    except ValueError:
        return tanggal_str

# Fungsi untuk konversi DOCX ke PDF menggunakan LibreOffice
def convert_docx_to_pdf(docx_path):
    if not os.path.isfile(docx_path):
        raise FileNotFoundError(f"File tidak ditemukan: {docx_path}")

    output_dir = os.path.dirname(docx_path)
    try:
        subprocess.run([
            "libreoffice", "--headless", "--convert-to", "pdf", docx_path, "--outdir", output_dir
        ], check=True)
    except subprocess.CalledProcessError as e:
        return f"Error saat mengonversi ke PDF: {e}"

    # Path PDF yang dihasilkan
    pdf_path = os.path.splitext(docx_path)[0] + ".pdf"
    return pdf_path if os.path.exists(pdf_path) else None

@app.route('/')
def home():
    return render_template("index.html")

@app.route('/generate', methods=['POST'])
def generate():
    nama_petugas = request.form['nama_petugas'].strip().replace(" ", "_")
    timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
    template_path = os.path.join(TEMPLATE_FOLDER, 'template.docx')

    if not os.path.exists(template_path):
        return "Template tidak ditemukan."

    doc = Document(template_path)
    
    # Format tanggal surat
    tanggal_surat = request.form['tanggal_surat']
    lokasi = request.form.get('lokasi', '')

    # Tambahkan tanda tangan
    doc.add_paragraph(f"{lokasi}, {format_tanggal(tanggal_surat)}").alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    change_font(doc)

    # Simpan dokumen sementara
    docx_filename = f"laporan_{nama_petugas}_{timestamp}.docx"
    docx_path = os.path.join("temp", docx_filename)
    os.makedirs("temp", exist_ok=True)
    doc.save(docx_path)

    # Jika pengguna meminta file PDF
    if 'generate_pdf' in request.form:
        pdf_path = convert_docx_to_pdf(docx_path)
        if pdf_path:
            return send_file(pdf_path, as_attachment=True, download_name=f"Laporan_{nama_petugas}_{timestamp}.pdf", mimetype="application/pdf")
        else:
            return "Gagal mengonversi ke PDF.", 500

    # Jika hanya ingin file DOCX
    return send_file(docx_path, as_attachment=True, download_name=docx_filename, mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

if __name__ == '__main__':
    app.run(debug=True, host='0.0.0.0', port=5000)
